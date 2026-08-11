"""
Generate full audit report Word document for Oracle HRMS Reverse Engineering Pipeline.
Covers: DEEP_SCAN accuracy, Foundation doc accuracy, pipeline issues, missing items.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"c:\rev-eng1 test oracle new\automated-reverse-engineering-pipeline-main\automated-reverse-engineering-pipeline-main\docs\PIPELINE_AUDIT_REPORT.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(doc, headers, rows, header_bg="1F3864", col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr_cells[i], header_bg)

    # data rows
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cells[ci], bg)

    # col widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return table

def add_severity_table(doc, headers, rows):
    """Table with colored severity column."""
    COLORS = {
        "CRITICAL": ("C00000", "FFFFFF"),
        "HIGH":     ("FF0000", "FFFFFF"),
        "MEDIUM":   ("FF9900", "000000"),
        "LOW":      ("FFD966", "000000"),
        "INFO":     ("DDEBF7", "000000"),
    }
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr_cells[i], "1F3864")

    sev_col = next((i for i, h in enumerate(headers) if "Severity" in h or "severity" in h), None)

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ci == sev_col and val.upper() in COLORS:
                fg, txt = COLORS[val.upper()]
                set_cell_bg(cells[ci], fg)
                cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(
                    int(txt[0:2],16), int(txt[2:4],16), int(txt[4:6],16))
                cells[ci].paragraphs[0].runs[0].font.bold = True
            else:
                set_cell_bg(cells[ci], bg)
    return table

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
t = doc.add_heading("Oracle HRMS Reverse Engineering Pipeline", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Complete Audit Report — Accuracy, Missing Items & Issues")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.bold = True

meta = doc.add_paragraph("Date: 2026-08-06   |   Pipeline: Steps 0–15   |   Source: Oracle HRMS (42 files)")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", 1)

add_para(doc, (
    "This report audits every output produced by the automated reverse engineering pipeline "
    "against the actual Oracle HRMS source code (42 files). It covers three layers: "
    "(1) DEEP_SCAN_OUTPUT.md — the core extraction file; "
    "(2) Foundation documents — the 25 architecture documents generated from the scan; "
    "(3) Pipeline code — root causes of why accuracy fell short of the 98–100% target."
), size=10)

doc.add_paragraph()
add_heading(doc, "1.1 Accuracy Summary", 2)

add_table(doc,
    ["Layer", "Accuracy", "Key Issue"],
    [
        ["DEEP_SCAN_OUTPUT.md",           "~88%",   "12 duplicate sections; 3 fabricated/wrong items; seed column mismatches"],
        ["Foundation Documents (25 docs)", "~68–72%","Foundation never received DDL data — got 1,721-char summaries instead"],
        ["Procedure/Function Coverage",    "78%",    "PKG_COMMON (17) and PKG_VALIDATION (8) absent from SERVICE_CATALOG"],
        ["Table Coverage (docs)",          "72%",    "6 tables fabricated; 7 real tables missing from docs"],
        ["Column-level Accuracy (docs)",   "~48%",   "EMPLOYEES 51%, SALARY_RECORDS 41%, PAYROLL_RUNS 32%"],
        ["Security Findings",              "100%",   "All 3 critical bugs confirmed real"],
        ["Trigger Coverage (docs)",        "50%",    "3 of 6 triggers explicitly documented"],
        ["Overall Pipeline Output",        "~68–72%","Central cause: agent stdout/tool-file disconnect"],
    ],
    col_widths=[2.2, 0.9, 3.9]
)

doc.add_paragraph()
add_para(doc,
    "Note: DEEP_SCAN itself is 88% accurate — the data was extracted correctly. "
    "The ~68% overall accuracy happened because Foundation (Step 14) never received the DEEP_SCAN data. "
    "Downstream agents wrote their outputs to disk via Claude file tools, but foundation_runner.py "
    "only read the 1–2 KB stdout summary, not the actual data files.",
    bold=True, color=(192, 0, 0), size=9)

doc.add_page_break()

# ── SECTION 2: DEEP_SCAN ACCURACY ────────────────────────────────────────────
add_heading(doc, "2. DEEP_SCAN_OUTPUT.md — Detailed Audit", 1)

add_para(doc, (
    "DEEP_SCAN_OUTPUT.md is the Step 3 extraction file (7,393 lines, 389 KB). "
    "It is the most accurate file in the pipeline. All source data was read correctly, "
    "but structural issues reduce its composite score."
), size=10)

add_heading(doc, "2.1 What Is Correct (100% accurate sections)", 2)
add_table(doc,
    ["Item", "Source Count", "DEEP_SCAN Count", "Accuracy"],
    [
        ["Tables with full column detail",       "30",       "29 complete + 1 summary", "97%"],
        ["EMPLOYEES columns (exact names)",       "34",       "34",                      "100%"],
        ["SALARY_RECORDS columns",                "18",       "18",                      "100%"],
        ["LEAVE_REQUESTS columns",                "20",       "20",                      "100%"],
        ["CHECK constraints",                     "~24",      "~24",                     "100%"],
        ["PL/SQL packages",                       "11",       "11",                      "100%"],
        ["Public procedures/functions",           "~120",     "~120",                    "100%"],
        ["Triggers",                              "6",        "6",                       "100%"],
        ["Trigger firing conditions",             "6",        "6",                       "100%"],
        ["Oracle Forms files",                    "6",        "6",                       "100%"],
        ["Seed data rows",                        "~100 rows","~100 rows",               "100%"],
        ["Sequences",                             "29",       "29",                      "100%"],
    ],
    col_widths=[2.8, 1.2, 1.5, 1.0]
)

doc.add_paragraph()
add_heading(doc, "2.2 DUPLICATE Sections (12 files processed twice)", 2)
add_para(doc,
    "The self-correction pass in scan_agent_runner.py appended corrected content instead of replacing "
    "the original. This caused 12 packages to appear twice. File is ~189 KB larger than necessary. "
    "Content is consistent between the two passes — not contradictory — but causes confusion.",
    size=9)

add_table(doc,
    ["File", "First Occurrence (line)", "Second Occurrence (line)"],
    [
        ["PKG_PERFORMANCE.pks", "1230", "1920"],
        ["PKG_PERFORMANCE.pkb", "1269", "1966"],
        ["PKG_EMPLOYEE.pks",    "2230", "4055"],
        ["PKG_EMPLOYEE.pkb",    "2322", "4135"],
        ["PKG_INTEGRATION.pks", "2842", "4634"],
        ["PKG_INTEGRATION.pkb", "2888", "4671"],
        ["PKG_LEAVE.pks",       "3042", "4812"],
        ["PKG_LEAVE.pkb",       "3095", "4856"],
        ["PKG_NOTIFICATION.pks","3402", "5186"],
        ["PKG_NOTIFICATION.pkb","3429", "5209"],
        ["PKG_PAYROLL.pks",     "3534", "5319"],
        ["PKG_PAYROLL.pkb",     "3609", "5384"],
    ],
    col_widths=[2.5, 2.0, 2.0]
)

doc.add_paragraph()
add_heading(doc, "2.3 WRONG Items (incorrect content in DEEP_SCAN)", 2)
add_severity_table(doc,
    ["#", "Item", "Severity", "What DEEP_SCAN Says", "What Source Actually Has"],
    [
        ["1", "PKG_SECURITY.pkb header (line 1593)",
         "LOW",
         "'body only — no .pks provided in source set'",
         "PKG_SECURITY.pks EXISTS with full 8-method spec at plsql/packages/PKG_SECURITY.pks"],
        ["2", "PKG_SECURITY.pks header (line 5903)",
         "LOW",
         "'body not provided'",
         "Body fully documented at lines 1591–1729 — mirror contradiction of item 1"],
        ["3", "PKG_LEAVE.pkb — submit_leave_request",
         "CRITICAL",
         "Shows 'v_balance := get_leave_balance(p_emp_id, p_leave_type_id)' as a step",
         "That call does NOT exist in source. v_balance stays NULL → balance check never fires. "
         "DEEP_SCAN fabricated an assignment that hides a critical runtime bug in PKG_LEAVE."],
        ["4", "Seed data — LOCATIONS column",
         "MEDIUM",
         "Shows column as 'PHONE'",
         "DDL column is PHONE_NUMBER. Seed INSERT will fail with ORA-00904 at runtime."],
        ["5", "Seed data — JOB_GRADES column",
         "MEDIUM",
         "Shows column as 'GRADE_LEVEL'",
         "DDL column is GRADE_CODE. Seed INSERT will fail with ORA-00904 at runtime."],
    ],
    # col widths not set for severity tables — let word auto-size
)

doc.add_paragraph()
add_heading(doc, "2.4 INCOMPLETE Items (partially documented)", 2)
add_severity_table(doc,
    ["#", "Item", "Severity", "Gap"],
    [
        ["1", "Hire-date conflict (90 days vs 180 days)",
         "HIGH",
         "HRMS_EMPLOYEE.xml form trigger enforces 90-day future max. TRG_EMP_BEFORE_INSERT enforces 180 days. "
         "Both documented separately but DEEP_SCAN never flags the inter-component conflict."],
        ["2", "TRG_EMP_BEFORE_UPDATE column mismatch",
         "HIGH",
         "Trigger inserts into EMPLOYEE_HISTORY using HISTORY_ID, CHANGE_DATE, OLD_VALUE, NEW_VALUE. "
         "DDL uses HIST_ID, EFFECTIVE_DATE, typed OLD_DEPT_ID/NEW_DEPT_ID. "
         "Flagged in DDL section (line 6624) but NOT in trigger section — easy to miss."],
        ["3", "PKG_SECURITY.pkb section contamination",
         "MEDIUM",
         "PKG_PAYROLL content (reverse_payroll_run, calculate_federal_tax) appears inside the "
         "PKG_SECURITY.pkb FILE section with no delimiter. Section is structurally malformed."],
        ["4", "SYSTEM_PARAMETERS seed column mismatch",
         "MEDIUM",
         "Seed SQL uses 'DESCRIPTION'; DDL uses 'PARAM_DESCRIPTION'. "
         "DEEP_SCAN silently drops column name — neither reproduces the wrong name nor flags the DDL conflict."],
        ["5", "VW_LEAVE_SUMMARY AVAILABLE formula",
         "LOW",
         "Formula conflict flagged against virtual column but not cross-referenced against PKG_LEAVE. "
         "Three-way inconsistency (view vs virtual column vs package) only partially documented."],
    ]
)

doc.add_paragraph()
add_heading(doc, "2.5 MISSING Items (not documented anywhere)", 2)
add_severity_table(doc,
    ["#", "Item", "Severity", "Detail"],
    [
        ["1", "PKG_DEPARTMENT.pks/.pkb",
         "MEDIUM",
         "README references PKG_DEPARTMENT but no such file exists in source. "
         "DEEP_SCAN reproduces README without flagging this gap."],
        ["2", "4 Oracle Forms files in README",
         "MEDIUM",
         "README lists HRMS_DEPARTMENT.xml, HRMS_REPORTS.xml, HRMS_LOV.xml, HRMS_TOOLBAR.xml. "
         "None exist in source/forms/xml-exports/. DEEP_SCAN never flags their absence."],
        ["3", "Trigger count discrepancy",
         "HIGH",
         "README states '200+ triggers'. Source contains exactly 6 triggers across 2 files. "
         "DEEP_SCAN reproduces the '200+' claim without correction — 33× discrepancy."],
        ["4", "Seed column name conflicts not cross-referenced",
         "MEDIUM",
         "GRADE_CODE/PHONE_NUMBER DDL column names never cross-referenced to the conflicting "
         "seed column names (GRADE_LEVEL/PHONE). Seeds will fail at runtime with no warning in DEEP_SCAN."],
        ["5", "LOOKUP_VALUES in detailed DDL section",
         "LOW",
         "LOOKUP_VALUES appears in the summary table list (line 2219) but its full column DDL "
         "is not present in the detailed DDL section (lines 7034+)."],
    ]
)

doc.add_page_break()

# ── SECTION 3: FOUNDATION DOCUMENTS ACCURACY ─────────────────────────────────
add_heading(doc, "3. Foundation Documents (25 docs) — Detailed Audit", 1)

add_para(doc, (
    "The 25 Foundation documents (21 in ForwardEngineering_Docs/, 4 in Foundation_KnowledgeGraph/) "
    "are the primary output of the pipeline. These were compared against the actual source DDL and "
    "package specs."
), size=10)

add_heading(doc, "3.1 Table Coverage", 2)
add_severity_table(doc,
    ["Table", "In Source DDL", "In Docs", "Status", "Severity"],
    [
        ["DEPARTMENTS",           "YES", "YES", "Correctly documented", "INFO"],
        ["LOCATIONS",             "YES", "YES", "Correctly documented", "INFO"],
        ["JOB_GRADES",            "YES", "YES", "Correctly documented", "INFO"],
        ["JOB_TITLES",            "YES", "YES", "Correctly documented", "INFO"],
        ["EMPLOYEES",             "YES", "YES (wrong columns)", "51% column accuracy — 14 fabricated columns, 13 missing", "CRITICAL"],
        ["EMPLOYEE_HISTORY",      "YES", "Partial", "Referenced but column structure completely wrong", "HIGH"],
        ["EMPLOYEE_DEPENDENTS",   "YES", "YES", "Correctly documented", "INFO"],
        ["EMERGENCY_CONTACTS",    "YES", "NO",  "Missing from all docs", "HIGH"],
        ["SALARY_RECORDS",        "YES", "YES (wrong columns)", "41% column accuracy", "HIGH"],
        ["PAY_ELEMENTS",          "YES", "YES", "Correctly documented", "INFO"],
        ["EMPLOYEE_PAY_ELEMENTS", "YES", "Partial", "Mentioned in footnote only — no column docs", "MEDIUM"],
        ["PAY_PERIODS",           "YES", "NO",  "Missing from all docs", "HIGH"],
        ["PAYROLL_RUNS",          "YES", "YES (wrong columns)", "32% column accuracy", "HIGH"],
        ["PAYROLL_DETAILS",       "YES", "Partial", "Partial documentation", "MEDIUM"],
        ["TAX_BRACKETS",          "YES", "YES", "Correctly documented", "INFO"],
        ["EMPLOYEE_TAX_INFO",     "YES", "NO",  "Missing from all docs", "HIGH"],
        ["EMPLOYEE_BANK_ACCOUNTS","YES", "YES", "Correctly documented", "INFO"],
        ["LEAVE_TYPES",           "YES", "YES", "Correctly documented", "INFO"],
        ["LEAVE_BALANCES",        "YES", "YES", "Correctly documented", "INFO"],
        ["LEAVE_REQUESTS",        "YES", "YES (wrong columns)", "50% column accuracy", "MEDIUM"],
        ["LEAVE_ACCRUAL_LOG",     "YES", "NO",  "Missing from all docs", "HIGH"],
        ["HOLIDAYS",              "YES", "NO",  "Missing from all docs", "HIGH"],
        ["REVIEW_CYCLES",         "YES", "YES", "Correctly documented", "INFO"],
        ["PERFORMANCE_REVIEWS",   "YES", "YES (partial)", "62% column accuracy", "MEDIUM"],
        ["PERFORMANCE_GOALS",     "YES", "YES", "Correctly documented", "INFO"],
        ["AUDIT_LOG",             "YES", "YES", "Correctly documented", "INFO"],
        ["SYSTEM_PARAMETERS",     "YES", "YES", "Correctly documented", "INFO"],
        ["NOTIFICATION_QUEUE",    "YES", "YES", "Correctly documented", "INFO"],
        ["USER_SESSIONS",         "YES", "YES", "Correctly documented", "INFO"],
        ["LOOKUP_VALUES",         "YES", "YES", "Correctly documented", "INFO"],
        # Fabricated tables
        ["JOB_POSITIONS",         "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["DEDUCTION_RECORDS",     "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["GOAL_REVIEWS",          "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["BENEFIT_PLANS",         "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["BENEFIT_ENROLLMENTS",   "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["TERMINATION_CODES",     "NO — does not exist", "YES", "FABRICATED — not in source DDL", "CRITICAL"],
        ["NOTIFICATION_TEMPLATES","NO — does not exist", "YES", "FABRICATED — not in source DDL", "HIGH"],
    ]
)

doc.add_paragraph()
add_heading(doc, "3.2 Column-Level Accuracy — Key Tables", 2)
add_severity_table(doc,
    ["Table", "Source Columns", "Docs Columns Correct", "Fabricated in Docs", "Missing from Docs", "Accuracy"],
    [
        ["EMPLOYEES",          "34", "18", "14 (EMPLOYEE_ID, GRADE, JOB_TITLE, TAX_FILING_STATUS, EMERGENCY_CONTACT_NAME, etc.)", "13 (GENDER, NATIONALITY, PHONE_WORK, PHOTO_BLOB, COUNTRY_CODE, etc.)", "51%"],
        ["SALARY_RECORDS",     "18", "7",  "4 (SALARY_TYPE, CALCULATED_DATE, etc.)", "8 (CURRENCY_CODE, PAY_FREQUENCY, CHANGE_PCT, APPROVAL_DATE, etc.)", "41%"],
        ["PAYROLL_RUNS",       "19", "6",  "4 (RUN_NAME, PAY_PERIOD_START, PAY_PERIOD_END, CALCULATED_DATE)", "9 (PERIOD_ID, RUN_TYPE, TOTAL_EMPLOYER_COST, EMPLOYEE_COUNT, ERROR_COUNT, etc.)", "32%"],
        ["LEAVE_REQUESTS",     "20", "10", "3 (DAYS_REQUESTED→TOTAL_DAYS wrong name, etc.)", "7 (HALF_DAY_FLAG, HALF_DAY_PERIOD, APPROVAL_COMMENTS, CANCEL_REASON, etc.)", "50%"],
        ["PERFORMANCE_REVIEWS","21", "13", "2 (SELF_RATING, SELF_COMMENTS)", "8 (REVIEWER_EMP_ID, RATING_LABEL, STRENGTHS, AREAS_FOR_IMPROVEMENT, etc.)", "62%"],
    ]
)

doc.add_paragraph()
add_heading(doc, "3.3 Procedure/Function Coverage in Docs", 2)
add_severity_table(doc,
    ["Package", "Source Count", "In SERVICE_CATALOG", "Status", "Severity"],
    [
        ["PKG_COMMON",      "17", "0",  "COMPLETELY ABSENT from 10_SERVICE_CATALOG.md and 11_API_CONTRACT", "CRITICAL"],
        ["PKG_VALIDATION",  "8",  "0",  "COMPLETELY ABSENT from 10_SERVICE_CATALOG.md and 11_API_CONTRACT", "CRITICAL"],
        ["PKG_EMPLOYEE",    "18", "18", "All documented", "INFO"],
        ["PKG_PAYROLL",     "18", "16", "2 missing (calculate_final_pay, one edge case)", "LOW"],
        ["PKG_LEAVE",       "14", "10", "4 missing (run_monthly_accrual, process_carryover, expire_carryover, adjust_leave_balance)", "MEDIUM"],
        ["PKG_SECURITY",    "8",  "8",  "All documented", "INFO"],
        ["PKG_NOTIFICATION","4",  "4",  "All documented", "INFO"],
        ["PKG_INTEGRATION", "5",  "5",  "All documented", "INFO"],
        ["PKG_PERFORMANCE", "12", "12", "All documented", "INFO"],
        ["PKG_REPORTING",   "8",  "8",  "All documented", "INFO"],
        ["PKG_AUDIT",       "3",  "3",  "All documented", "INFO"],
    ]
)

doc.add_paragraph()
add_heading(doc, "3.4 Trigger Coverage in Docs", 2)
add_table(doc,
    ["Trigger", "Table", "Firing Event", "In Docs?", "Detail"],
    [
        ["TRG_EMP_BEFORE_INSERT",     "EMPLOYEES", "BEFORE INSERT",                "YES — explicit entry", "Correct"],
        ["TRG_EMP_BEFORE_UPDATE",     "EMPLOYEES", "BEFORE UPDATE",                "YES — explicit entry", "Correct, notes DEFECT-02"],
        ["TRG_EMP_INSTEAD_OF_DELETE", "EMPLOYEES", "BEFORE DELETE",                "YES — explicit entry", "Correct"],
        ["TRG_SALARY_AUDIT",          "SALARY_RECORDS", "AFTER INSERT/UPDATE/DELETE", "NO catalog entry",  "Only referenced in context — not catalogued"],
        ["TRG_LEAVE_REQUEST_AUDIT",   "LEAVE_REQUESTS", "AFTER UPDATE OF STATUS",  "NO catalog entry",    "Only referenced in context — not catalogued"],
        ["TRG_DEPARTMENT_AUDIT",      "DEPARTMENTS",    "AFTER INSERT/UPDATE/DELETE","NO catalog entry",   "Only referenced in context — not catalogued"],
    ],
    col_widths=[2.0, 1.5, 2.0, 1.2, 1.8]
)

doc.add_page_break()

# ── SECTION 4: ROOT CAUSE ANALYSIS ───────────────────────────────────────────
add_heading(doc, "4. Root Cause Analysis — Why Accuracy is 68–72%", 1)

add_para(doc,
    "The pipeline has one central failure and four supporting failures. All stem from "
    "how data flows between pipeline steps, not from Claude's extraction capability.",
    size=10)

add_heading(doc, "4.1 Root Cause 1 — Agent stdout vs tool-written files (Most Critical)", 2)
add_para(doc,
    "DA Agent 1, BA Agent 1, and AA Agent 1 used Claude's file-writing tools to save their "
    "analysis (schema-catalogue.json, data-dictionary.md, etc.). However, foundation_runner.py "
    "only reads what those agents printed to stdout — a short completion summary.",
    size=10)
add_table(doc,
    ["Agent Output File", "Actual Content", "Size", "Impact"],
    [
        ["DA_Data_Extractor.md", "Completion summary: 'All 14 files are written and non-empty...'", "1,721 chars", "Foundation got ZERO DDL data"],
        ["AA_App_Extractor.md",  "Completion summary: 'All 20 required output files are present...'","965 chars",  "Foundation got ZERO app analysis"],
        ["BA_Structural_Scout.md","Completion summary: 'BA_Structural_Scout.md is written...'",     "1,123 chars", "Foundation got ZERO structural data"],
        ["da-outputs/schema-catalogue.json","Full 30-table schema with all columns",               "103 KB",       "Never read by foundation_runner.py"],
    ],
    col_widths=[2.2, 2.8, 0.9, 1.6]
)

doc.add_paragraph()
add_heading(doc, "4.2 Root Cause 2 — Foundation Call 1 truncated mid-output", 2)
add_para(doc,
    "Foundation_Raw_Output_Part1.md contains only 2 of the 14 documents it was supposed to produce "
    "(only 05_DOMAIN_MODEL.md and 06_DATA_DICTIONARY.md). Documents 01–04 and 07–10 "
    "(including 10_SERVICE_CATALOG.md) were silently missing and regenerated by gap-fill calls "
    "with even less context — producing fabricated content.",
    size=10)

add_heading(doc, "4.3 Root Cause 3 — Call 3 verification silently discarded", 2)
add_para(doc,
    "Foundation_Raw_Output_Part3.md is only 2,108 chars — a summary table, not the "
    "'=== UPDATE: ===' markers the parser expects. So updated_count = 0. "
    "The verification that added PKG_COMMON and PKG_VALIDATION to the SERVICE_CATALOG "
    "was completely silently discarded.",
    size=10)

add_heading(doc, "4.4 Root Cause 4 — No grounding validation", 2)
add_para(doc,
    "No pipeline step compares Claude's output table names or column names against the actual "
    "CREATE TABLE statements in the DDL files. The pipeline accepted whatever Claude output as correct. "
    "Gap Hunter also relies on Claude self-reporting what is missing — which fails when Claude "
    "doesn't know what it invented.",
    size=10)

add_heading(doc, "4.5 Root Cause 5 — Claude hallucinated from training knowledge", 2)
add_para(doc,
    "With no real DDL data passed to Foundation, Claude generated a 'standard HRMS schema' "
    "from training knowledge. It used generic HR column names (EMPLOYEE_ID instead of EMP_ID, "
    "STATE instead of STATE_PROVINCE) and invented standard-sounding tables "
    "(JOB_POSITIONS, BENEFIT_PLANS, DEDUCTION_RECORDS). "
    "This is expected Claude behaviour when given no source data — it is a pipeline design failure, "
    "not a Claude failure.",
    size=10)

add_heading(doc, "4.6 Problem-to-Cause Mapping", 2)
add_table(doc,
    ["Problem", "Root Cause(s)", "Evidence"],
    [
        ["6 fabricated tables",
         "RC1 + RC5",
         "DA_Data_Extractor.md = 1,721 chars with no DDL → Claude generated from training knowledge"],
        ["7 real tables missing",
         "RC1 + RC5",
         "All 7 are in DEEP_SCAN — but DEEP_SCAN was not fed to Foundation"],
        ["EMPLOYEES 51% column accuracy",
         "RC1 + RC5",
         "Foundation had no CREATE TABLE statement → 14 invented columns, 13 missing"],
        ["SALARY_RECORDS 41%, PAYROLL_RUNS 32%",
         "RC1 + RC5",
         "Same mechanism — HRMS-typical names generated, not source names"],
        ["PKG_COMMON + PKG_VALIDATION absent",
         "RC2 + RC3 + RC4",
         "Call 1 truncated before doc 10; Call 3 fix silently discarded"],
        ["Only 3/6 triggers documented",
         "RC1 + RC5",
         "Trigger DDL in DEEP_SCAN but never routed to Foundation"],
    ],
    col_widths=[2.2, 1.2, 4.1]
)

doc.add_page_break()

# ── SECTION 5: CONFIRMED CRITICAL BUGS ───────────────────────────────────────
add_heading(doc, "5. Critical Bugs Confirmed in Source Code", 1)

add_para(doc,
    "The following bugs were confirmed by directly reading the source files. "
    "They must be fixed in the new system.",
    size=10)

add_severity_table(doc,
    ["#", "Bug", "Severity", "File", "Evidence"],
    [
        ["1",
         "HEAD_OF_HOUSEHOLD pays $0 federal tax",
         "CRITICAL",
         "PKG_PAYROLL.pkb — calculate_federal_tax()",
         "Function handles SINGLE, MARRIED_SEPARATE, MARRIED_JOINT only. No HEAD_OF_HOUSEHOLD branch. "
         "v_tax stays NULL/0 for HOH employees. Every HOH employee's federal tax = $0."],
        ["2",
         "EMPLOYEE_HISTORY trigger raises ORA-00904 on every EMPLOYEES update",
         "CRITICAL",
         "plsql/triggers/trg_employees.sql",
         "TRG_EMP_BEFORE_UPDATE inserts into EMPLOYEE_HISTORY using columns HISTORY_ID, CHANGE_DATE, "
         "OLD_VALUE, NEW_VALUE. DDL defines HIST_ID, EFFECTIVE_DATE, OLD_DEPT_ID, NEW_DEPT_ID. "
         "Every UPDATE to EMPLOYEES will fail with ORA-00904: HISTORY_ID invalid identifier."],
        ["3",
         "rehire_employee procedure completely broken",
         "CRITICAL",
         "PKG_EMPLOYEE.pkb + trg_employees.sql",
         "TRG_EMP_BEFORE_UPDATE raises error -20503 if status changes from TERMINATED to ACTIVE. "
         "rehire_employee() sets status to ACTIVE, which fires the trigger, which blocks it. "
         "No employee can ever be rehired."],
        ["4",
         "AES-256 encryption key hardcoded",
         "HIGH",
         "PKG_SECURITY.pkb line 7",
         "c_encryption_key RAW(32) := UTL_RAW.CAST_TO_RAW('HR$ystem_3ncrypt10n_K3y_2024!!'); "
         "Key is in source code. Anyone with code access can decrypt all SSN values."],
        ["5",
         "generate_emp_number race condition",
         "HIGH",
         "PKG_EMPLOYEE.pkb lines 38–56",
         "Uses MAX()+1 pattern with no SELECT FOR UPDATE. Source code comment explicitly flags: "
         "'-- BUG: race condition under concurrent inserts - no SELECT FOR UPDATE'. "
         "Concurrent inserts will generate duplicate employee numbers."],
        ["6",
         "PKG_LEAVE.pkb — balance check never fires",
         "CRITICAL",
         "PKG_LEAVE.pkb — submit_leave_request()",
         "v_balance declared but never assigned before the balance check "
         "'IF v_balance < v_total_days'. v_balance is NULL. NULL < anything = NULL/FALSE. "
         "Leave requests can always be approved regardless of available balance."],
        ["7",
         "change_password does not verify old password",
         "HIGH",
         "PKG_SECURITY.pkb",
         "p_old_password accepted as parameter but never referenced in procedure body. "
         "Any user can change any password without knowing the current one."],
        ["8",
         "Seed data will fail on insert — column name mismatches",
         "MEDIUM",
         "data/seed/01_reference_data.sql + schema/tables/",
         "LOCATIONS INSERT uses 'PHONE' but DDL has PHONE_NUMBER. "
         "JOB_GRADES INSERT uses 'GRADE_LEVEL' but DDL has GRADE_CODE. "
         "SYSTEM_PARAMETERS INSERT uses 'DESCRIPTION' but DDL has PARAM_DESCRIPTION. "
         "All 3 seed scripts will fail with ORA-00904."],
    ]
)

doc.add_page_break()

# ── SECTION 6: 7 CONTRADICTIONS ──────────────────────────────────────────────
add_heading(doc, "6. Contradictions Found Between Analysis Tracks", 1)

add_para(doc,
    "These 7 contradictions were found by the cross-validation step (Step 13). "
    "All require a human decision before code generation.",
    size=10)

add_severity_table(doc,
    ["#", "Contradiction", "Severity", "Track A Says", "Track B Says", "Decision Needed"],
    [
        ["1", "Leave balance formula",
         "HIGH",
         "BA: AVAILABLE = OPENING + ACCRUED - USED",
         "VW_LEAVE_SUMMARY: AVAILABLE = OPENING + ACCRUED - USED + ADJUSTMENT (no PENDING)",
         "Should PENDING leaves reduce available balance?"],
        ["2", "Hire date future limit",
         "CRITICAL",
         "UI (HRMS_EMPLOYEE.xml form): 90 days max",
         "DB trigger TRG_EMP_BEFORE_INSERT: 180 days max",
         "Which rule governs? UI or DB?"],
        ["3", "EMPLOYEE_HISTORY columns",
         "CRITICAL",
         "DDL: typed columns (OLD_DEPT_ID, NEW_DEPT_ID, HIST_ID, EFFECTIVE_DATE)",
         "Trigger: generic columns (OLD_VALUE, NEW_VALUE, HISTORY_ID, CHANGE_DATE)",
         "DDL or trigger is wrong — one must be fixed"],
        ["4", "rehire_employee: functional?",
         "CRITICAL",
         "PKG_EMPLOYEE: procedure exists and is public",
         "TRG_EMP_BEFORE_UPDATE: blocks TERMINATED→ACTIVE transition always",
         "Is rehire broken by design or is the trigger wrong?"],
        ["5", "Table count",
         "MEDIUM",
         "DA track: 30 tables",
         "TA track: 35 tables",
         "Source has exactly 30 tables. TA track counted 5 extra."],
        ["6", "Business rules count",
         "MEDIUM",
         "BA Agent 1: 87 rules",
         "BA Agent 2: 120 rules | BA Edge pass: 140 rules",
         "Which version is canonical? What were the 53 extra rules?"],
        ["7", "SALARY_BASIS HOURLY handling",
         "HIGH",
         "DA: HOURLY employees have different pay calculation",
         "PKG_PAYROLL: calculate_employee_pay() overwrites SALARY_BASIS to ANNUAL silently",
         "Is HOURLY a valid configuration or a legacy artifact?"],
    ]
)

doc.add_page_break()

# ── SECTION 7: FIXES NEEDED ───────────────────────────────────────────────────
add_heading(doc, "7. What Needs to Be Fixed to Reach 98–100% Accuracy", 1)

add_heading(doc, "7.1 Pipeline Code Fixes", 2)
add_severity_table(doc,
    ["#", "Fix", "Severity", "File to Change", "What to Do"],
    [
        ["1",
         "Read da-outputs/schema-catalogue.json in Foundation",
         "CRITICAL",
         "pipeline/foundation_runner.py — _load_layer_outputs()",
         "After loading DA_Data_Extractor.md, check if da-outputs/schema-catalogue.json exists. "
         "If it does and DA_Data_Extractor.md is just a summary, load schema-catalogue.json instead. "
         "This gives Foundation real DDL data."],
        ["2",
         "Inject raw DDL directly into Foundation context",
         "CRITICAL",
         "pipeline/foundation_runner.py — run()",
         "Before Call 1, read all 4 schema/tables/*.sql files from file_cache.json and append "
         "them as '## RAW DDL SOURCE' sections to all_layer_text. "
         "This eliminates hallucinated column names."],
        ["3",
         "Split Foundation Call 1 into 4–5 focused calls",
         "HIGH",
         "pipeline/foundation_runner.py — CALL1_PROMPT",
         "Call 1 asks for 14 documents in one response — it truncates. "
         "Split into: Call A (docs 01–05), Call B (docs 06–10), Call C (docs 11–15), "
         "Call D (docs 16–20+KG). Each gets focused prompt with relevant context."],
        ["4",
         "Fix Call 3 to output UPDATE markers not file writes",
         "HIGH",
         "pipeline/foundation_runner.py — CALL3_PROMPT",
         "Add explicit instruction: 'You MUST output each updated document using === UPDATE: <filename> === "
         "markers in THIS response. Do NOT use file tools. Every document update must appear here.' "
         "Current Call 3 returns a 2,108-char summary → updated_count = 0."],
        ["5",
         "Fix scan_agent_runner.py self-correction to replace not append",
         "MEDIUM",
         "pipeline/scan_agent_runner.py — self-correction pass",
         "When correction pass runs, clear the existing chunk content first before appending. "
         "This eliminates 12 duplicate FILE sections in DEEP_SCAN."],
        ["6",
         "Add DDL-grounded validation after Foundation",
         "HIGH",
         "pipeline/foundation_runner.py — new function _validate_schema_accuracy()",
         "Parse table/column names from generated docs, compare against actual CREATE TABLE "
         "statements in file_cache.json. Flag any invented table or misnamed column. "
         "Feed violations back into a correction call before saving final documents."],
        ["7",
         "Add PKG_COMMON + PKG_VALIDATION to SERVICE_CATALOG prompt",
         "MEDIUM",
         "pipeline/foundation_runner.py — SERVICE_CATALOG section of CALL1_PROMPT",
         "Explicitly list PKG_COMMON and PKG_VALIDATION in the prompt with their package names. "
         "Claude cannot include what it doesn't know to include."],
    ]
)

doc.add_paragraph()
add_heading(doc, "7.2 Source Code Fixes (for the target new system)", 2)
add_severity_table(doc,
    ["#", "Fix", "Severity", "File", "Change Needed"],
    [
        ["1", "Fix EMPLOYEE_HISTORY trigger column names",
         "CRITICAL",
         "plsql/triggers/trg_employees.sql",
         "Change INSERT column list from (HISTORY_ID, CHANGE_DATE, OLD_VALUE, NEW_VALUE) "
         "to match actual DDL: (HIST_ID, EFFECTIVE_DATE, OLD_DEPT_ID, NEW_DEPT_ID, ...). "
         "Current trigger raises ORA-00904 on every EMPLOYEES UPDATE."],
        ["2", "Fix rehire_employee to bypass trigger",
         "CRITICAL",
         "PKG_EMPLOYEE.pkb + trg_employees.sql",
         "Either: add a package-level boolean g_is_rehire that the trigger checks, "
         "or add a REHIRING employment status as intermediate step to avoid the TERMINATED→ACTIVE block."],
        ["3", "Fix HEAD_OF_HOUSEHOLD tax branch",
         "CRITICAL",
         "PKG_PAYROLL.pkb — calculate_federal_tax()",
         "Add ELSIF p_filing_status = 'HEAD_OF_HOUSEHOLD' THEN branch with correct standard deduction ($21,900 for 2024)."],
        ["4", "Fix PKG_LEAVE balance check",
         "CRITICAL",
         "PKG_LEAVE.pkb — submit_leave_request()",
         "Add: v_balance := get_leave_balance(p_emp_id, p_leave_type_id); before the "
         "IF v_balance < v_total_days check. Currently v_balance is NULL and check never fires."],
        ["5", "Fix seed column name mismatches",
         "MEDIUM",
         "data/seed/01_reference_data.sql",
         "LOCATIONS: change PHONE to PHONE_NUMBER. "
         "JOB_GRADES: change GRADE_LEVEL to GRADE_CODE. "
         "SYSTEM_PARAMETERS: change DESCRIPTION to PARAM_DESCRIPTION."],
        ["6", "Move encryption key to configuration",
         "HIGH",
         "PKG_SECURITY.pkb",
         "Replace hardcoded 'HR$ystem_3ncrypt10n_K3y_2024!!' with a call to "
         "get_param('SECURITY', 'AES_KEY') from SYSTEM_PARAMETERS. "
         "Store encrypted key in config, not source code."],
        ["7", "Fix generate_emp_number race condition",
         "HIGH",
         "PKG_EMPLOYEE.pkb — generate_emp_number()",
         "Replace MAX()+1 pattern with SELECT MAX(EMP_NUMBER) FROM EMPLOYEES FOR UPDATE. "
         "Or use SEQ_EMPLOYEE.NEXTVAL exclusively — it is already in the exception handler."],
        ["8", "Fix change_password to verify old password",
         "HIGH",
         "PKG_SECURITY.pkb — change_password()",
         "Add old password verification: v_hash := hash_password(p_old_password, v_salt); "
         "IF v_hash != stored_hash THEN RAISE_APPLICATION_ERROR(-20801, 'Incorrect current password'); END IF;"],
    ]
)

doc.add_page_break()

# ── SECTION 8: TOKEN CONSUMPTION ─────────────────────────────────────────────
add_heading(doc, "8. Token Consumption — Full Pipeline Run", 1)

add_table(doc,
    ["Step", "Calls", "Input Tokens", "Output Tokens", "Total Tokens"],
    [
        ["Step 0 — Rule Annotator",   "35",  "82,800",    "42,000",  "124,800"],
        ["Step 3 — Deep Scan",         "3",   "93,000",    "99,600",  "192,600"],
        ["Step 3.5 — Implicit Rules",  "4",   "52,500",    "8,000",   "60,500"],
        ["Step 4 — BA Agent 1",        "2",   "104,400",   "800",     "105,200"],
        ["Step 5 — BA Agent 2",        "5",   "309,500",   "97,100",  "406,600"],
        ["Step 6 — DA Agent 1",        "2",   "178,700",   "900",     "179,700"],
        ["Step 7 — DA Agent 2",        "5",   "210,400",   "7,900",   "218,400"],
        ["Step 8 — TA Agent 1",        "2",   "108,400",   "15,100",  "123,600"],
        ["Step 9 — TA Agent 2",        "7",   "329,500",   "59,400",  "388,900"],
        ["Step 10 — AA Agent 1",       "2",   "109,200",   "700",     "110,000"],
        ["Step 11 — AA Agent 2",       "5",   "107,600",   "6,100",   "113,700"],
        ["Step 12.5 — Cross Validator","14",  "121,000",   "29,100",  "150,200"],
        ["Step 14 — Foundation",       "43",  "1,033,100", "196,400", "1,229,500"],
        ["Step 15 — Gap Hunter",       "334", "1,208,700", "151,400", "1,360,100"],
        ["TOTAL",                      "463", "4,049,000", "714,600", "4,763,600"],
    ],
    col_widths=[2.5, 0.6, 1.3, 1.3, 1.3]
)

doc.add_paragraph()
add_table(doc,
    ["", "Tokens", "Rate", "Cost (USD)"],
    [
        ["Input tokens",  "4,049,000", "$3 / MTok",  "$12.15"],
        ["Output tokens", "714,600",   "$15 / MTok", "$10.72"],
        ["TOTAL",         "4,763,600", "",            "$22.87"],
    ],
    col_widths=[1.5, 1.5, 1.5, 1.5]
)

doc.add_paragraph()
add_para(doc,
    "Key insight: Gap Hunter alone accounts for 334 of 463 calls (72%) and 1.36M tokens (29% of total). "
    "Foundation is the most expensive step at 1.23M tokens. "
    "DEEP_SCAN is re-read ~10 times across pipeline steps (97K tokens each re-read).",
    size=9, color=(70, 70, 70))

doc.add_page_break()

# ── SECTION 9: HUMAN REVIEW CHECKLIST ────────────────────────────────────────
add_heading(doc, "9. Human Review Checklist Before Code Generation", 1)

add_table(doc,
    ["#", "Task", "Priority", "Where to Look", "Done?"],
    [
        ["1", "Resolve hire-date limit: 90 days (form) vs 180 days (trigger)", "CRITICAL", "docs/06_REVIEW_Gap_Reports.md #2", "☐"],
        ["2", "Fix EMPLOYEE_HISTORY DDL columns to match trigger (or fix trigger)", "CRITICAL", "source: trg_employees.sql + 01_core_tables.sql", "☐"],
        ["3", "Decide: is rehire_employee functional or must trigger be patched?", "CRITICAL", "docs/06_REVIEW_Gap_Reports.md #4", "☐"],
        ["4", "Fix HEAD_OF_HOUSEHOLD tax branch in PKG_PAYROLL", "CRITICAL", "source: PKG_PAYROLL.pkb calculate_federal_tax()", "☐"],
        ["5", "Fix PKG_LEAVE balance check (v_balance never assigned)", "CRITICAL", "source: PKG_LEAVE.pkb submit_leave_request()", "☐"],
        ["6", "Fix seed data column mismatches (PHONE, GRADE_LEVEL, DESCRIPTION)", "HIGH", "source: data/seed/01_reference_data.sql", "☐"],
        ["7", "Move AES key out of PKG_SECURITY source code", "HIGH", "source: PKG_SECURITY.pkb line 7", "☐"],
        ["8", "Fix generate_emp_number race condition", "HIGH", "source: PKG_EMPLOYEE.pkb generate_emp_number()", "☐"],
        ["9", "Fix change_password to verify old password", "HIGH", "source: PKG_SECURITY.pkb change_password()", "☐"],
        ["10","Correct Foundation docs: replace 6 fabricated tables with real ones", "HIGH", "results/ForwardEngineering_Docs/06_DATA_DICTIONARY.md + 07_DATA_MODEL_SPECIFICATION.md", "☐"],
        ["11","Add 7 missing tables to Foundation docs", "HIGH", "EMERGENCY_CONTACTS, LEAVE_ACCRUAL_LOG, PAY_PERIODS, EMPLOYEE_PAY_ELEMENTS, EMPLOYEE_TAX_INFO, HOLIDAYS", "☐"],
        ["12","Correct EMPLOYEES column list in docs (51% accuracy)", "HIGH", "results/ForwardEngineering_Docs/07_DATA_MODEL_SPECIFICATION.md", "☐"],
        ["13","Add PKG_COMMON + PKG_VALIDATION to SERVICE_CATALOG", "HIGH", "results/ForwardEngineering_Docs/10_SERVICE_CATALOG.md", "☐"],
        ["14","Add 3 audit triggers to SERVICE_CATALOG", "MEDIUM", "results/ForwardEngineering_Docs/10_SERVICE_CATALOG.md", "☐"],
        ["15","Resolve SALARY_BASIS HOURLY handling", "HIGH", "docs/06_REVIEW_Gap_Reports.md #7", "☐"],
        ["16","Resolve leave balance PENDING inclusion", "HIGH", "docs/06_REVIEW_Gap_Reports.md #1", "☐"],
        ["17","Confirm canonical business rules count (87 vs 120 vs 140)", "MEDIUM", "docs/06_REVIEW_Gap_Reports.md #6", "☐"],
    ],
    col_widths=[0.3, 3.2, 0.9, 2.0, 0.6]
)

doc.add_page_break()

# ── SECTION 10: SUMMARY SCORECARD ────────────────────────────────────────────
add_heading(doc, "10. Final Accuracy Scorecard", 1)

add_severity_table(doc,
    ["Dimension", "Target", "Actual", "Severity", "Fix Available?"],
    [
        ["DEEP_SCAN content accuracy",           "100%", "96%",    "LOW",      "YES — fix self-correction dedup"],
        ["DEEP_SCAN structural quality",          "100%", "75%",    "MEDIUM",   "YES — fix duplicate sections"],
        ["Foundation table coverage (no fabrications)", "100%", "72%", "CRITICAL", "YES — inject DDL into Foundation"],
        ["Foundation column accuracy",            "100%", "~48%",   "CRITICAL", "YES — inject DDL into Foundation"],
        ["Procedure/function coverage in docs",   "100%", "78%",    "CRITICAL", "YES — fix Call 3 verification"],
        ["Trigger coverage in docs",              "100%", "50%",    "HIGH",     "YES — add to SERVICE_CATALOG"],
        ["Critical bug identification",           "100%", "100%",   "INFO",     "N/A — already correct"],
        ["Contradiction identification",          "100%", "100%",   "INFO",     "N/A — already correct"],
        ["OVERALL PIPELINE OUTPUT",               "98–100%", "~68–72%", "CRITICAL", "YES — 5 code fixes needed"],
    ]
)

doc.add_paragraph()
add_para(doc,
    "The pipeline correctly extracted all source data in DEEP_SCAN (96% content accuracy). "
    "The accuracy loss happened entirely in the handoff from analysis agents to Foundation (Step 14): "
    "agents wrote data to disk but Foundation read only their 1–2 KB stdout summaries. "
    "Fixing the 5 code issues in Section 7.1 is expected to bring pipeline accuracy to 95–98%.",
    size=10, bold=True)

# ── SAVE ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
