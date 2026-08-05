"""
Word document: Full end-to-end explanation of the reverse engineering pipeline itself.
How it works, why it works that way, each step explained fully.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(11)
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def nb(text):
    p = doc.add_paragraph()
    r = p.add_run("NOTE: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0,0x50,0x00)
    p.add_run(text)
    return p

def diagram(text):
    p = doc.add_paragraph(text)
    p.style = 'No Spacing'
    for r in p.runs:
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x00,0x00,0x80)
    return p

def grid(headers, data):
    t = doc.add_table(rows=1+len(data), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for j,h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
    for i,row in enumerate(data):
        for j,val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
    doc.add_paragraph()

def info(rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Shading Accent 1'
    for i,(k,v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for r in t.rows[i].cells[0].paragraphs[0].runs: r.bold = True
    doc.add_paragraph()

# ══════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Oracle HRMS Reverse Engineering Pipeline")
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("End-to-End Working — How the Pipeline Works")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x2E,0x74,0xB5)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("A complete guide to the automated AI-powered reverse engineering system\nBuilt with Python + Claude AI (Anthropic) | 15 Steps | 4 Parallel Analysis Tracks")

doc.add_page_break()

# ══════════════════════════════════════════
# 1. WHAT IS THIS PROJECT
# ══════════════════════════════════════════
h1("1. What Is This Project?")

body(
    "This is an automated reverse engineering pipeline. It takes an existing Oracle HRMS "
    "system — which is written in PL/SQL and Oracle Forms — and fully analyses it using AI, "
    "producing complete documentation, business rules, data models, architecture diagrams, "
    "and migration plans — without any manual code reading."
)

body(
    "The problem it solves: most legacy Oracle systems were built 10–20 years ago. The original "
    "developers have left. No documentation exists. The system works — but nobody fully understands "
    "what it does, what business rules it enforces, or how its data is structured. Reverse engineering "
    "recovers all of this automatically."
)

h2("What Goes In, What Comes Out")
diagram("""
┌────────────────────────────────────────────────────────────────────┐
│                                                                    │
│   INPUT                   PIPELINE                  OUTPUT        │
│                                                                    │
│  source/            →   15-Step AI Pipeline   →   results/       │
│  ├── 22 PL/SQL           (Python + Claude AI)      ├── 140 Business Rules
│  │   packages                                      ├── Full Data Model (ERD)
│  ├── 6 Oracle Forms                                ├── Architecture Diagrams
│  │   XML exports                                   ├── Security Findings
│  ├── 11 SQL schema                                 ├── 25 Deliverable Docs
│  │   files                                         ├── Migration Plan
│  └── triggers,                                     └── Gap Report
│      views, seed data                                             │
│                                                                    │
│  42 files total                                 50+ output files  │
│  366,098 characters                             ~5 hours runtime  │
└────────────────────────────────────────────────────────────────────┘
""")

h2("The Core Idea: AI Reads the Code So You Don't Have To")
body(
    "Instead of a developer spending weeks reading thousands of lines of PL/SQL, "
    "the pipeline feeds all the source code to Claude AI in a structured way. "
    "Claude extracts business rules, maps data relationships, identifies bugs, "
    "finds security vulnerabilities, and writes professional deliverable documents — "
    "all automatically."
)

info([
    ("Language", "Python 3 (pipeline orchestration) + Claude AI (analysis)"),
    ("AI Engine", "Claude (Anthropic) called via the claude -p CLI command"),
    ("How AI is called", "Python subprocess — no API key setup needed, uses Claude CLI"),
    ("Total source input", "42 files, 366,098 characters of Oracle code"),
    ("Total runtime", "~5 hours end-to-end"),
    ("Output", "50+ files — analysis documents, JSON data, Markdown reports, architecture diagrams"),
    ("Resume-safe", "Yes — restart at any point, pipeline picks up exactly where it stopped"),
    ("Parallel processing", "4 analysis tracks run simultaneously to save time"),
])

doc.add_page_break()

# ══════════════════════════════════════════
# 2. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════
h1("2. Pipeline Architecture Overview")

h2("2.1 How Everything Connects")
diagram("""
┌─────────────────────────────────────────────────────────────────────┐
│                    PIPELINE ARCHITECTURE                            │
│                                                                     │
│  ┌──────────┐     ┌──────────────────────────────────────────────┐ │
│  │          │     │              run.py                          │ │
│  │ source/  │────▶│   The master orchestrator — calls each step  │ │
│  │ (42 files│     │   in the right order, manages parallel       │ │
│  │          │     │   threads, handles errors and retries        │ │
│  └──────────┘     └──────────┬───────────────────────────────────┘ │
│                              │                                      │
│                    ┌─────────▼─────────┐                           │
│                    │  pipeline/ folder │                           │
│                    │  One .py file     │                           │
│                    │  per step:        │                           │
│                    │                   │                           │
│  Step 1 ──────────▶  layer1_runner.py  │                           │
│  Step 2 ──────────▶  scan_once_runner  │                           │
│  Step 0 ──────────▶  rule_annotator_   │                           │
│  Step 3 ──────────▶  scan_runner.py    │                           │
│  Step 3.5 ────────▶  implicit_rules_   │                           │
│  Steps 4-12 ──────▶  ba/da/ta/aa_      │                           │
│  Step 13 ─────────▶  cross_validator_  │                           │
│  Step 14 ─────────▶  foundation_runner │                           │
│  Step 15 ─────────▶  gap_hunter_runner │                           │
│                    └─────────┬─────────┘                           │
│                              │                                      │
│                    ┌─────────▼─────────┐                           │
│                    │  base_runner.py   │                           │
│                    │                   │                           │
│                    │  call_claude()    │◀── Every step calls this  │
│                    │  • Runs: claude -p│    to invoke Claude AI    │
│                    │  • Retries 5x     │                           │
│                    │  • Returns text   │                           │
│                    └─────────┬─────────┘                           │
│                              │                                      │
│                    ┌─────────▼─────────┐                           │
│                    │  Claude AI (CLI)  │                           │
│                    │  claude -p        │                           │
│                    │  "prompt here"    │                           │
│                    └───────────────────┘                           │
│                                                                     │
│                    ┌───────────────────┐                           │
│                    │   results/        │                           │
│                    │   All outputs     │◀── Every step writes here │
│                    │   written here    │                           │
│                    └───────────────────┘                           │
└─────────────────────────────────────────────────────────────────────┘
""")

h2("2.2 Three Key Design Decisions")

h3("Decision 1: File Cache (file_cache.json)")
body(
    "In Step 2, the pipeline reads all 42 source files ONCE and stores them all in a single "
    "JSON file called file_cache.json. After that, the pipeline NEVER reads the source/ folder again."
)
body(
    "Every subsequent step reads from this cache. This means: (1) consistent data across all steps, "
    "(2) much faster — no disk reads after Step 2, (3) resume-safe — if a step fails, restarting "
    "reads from the same cache, not potentially changed source files."
)
diagram("""
WITHOUT cache:             WITH cache (this pipeline):
Step 3 reads source/  →    Step 2 reads source/ ONCE
Step 4 reads source/  →    file_cache.json created
Step 5 reads source/  →    Step 3 reads file_cache.json
Step 6 reads source/  →    Step 4 reads file_cache.json
(risk: files change        Step 5 reads file_cache.json
between steps)             (guaranteed consistent)
""")

h3("Decision 2: Parallel Analysis Tracks")
body(
    "Steps 4–12 are split into 4 tracks: Business Analysis (BA), Data Analysis (DA), "
    "Technology Analysis (TA), and Application Analysis (AA). These 4 tracks run "
    "simultaneously in Python threads — not one after another."
)
diagram("""
SEQUENTIAL (without parallelism):    PARALLEL (this pipeline):
BA track:  85 min  ─────────────────▶│
DA track:  48 min                    │ All 4 run at the SAME TIME
TA track:  86 min                    │ Wall clock = slowest track
AA track:  23 min                    │ = ~86 min total
                                     │ (not 242 min sequential)
TOTAL: 242 minutes                   TOTAL: ~86 minutes
""")

h3("Decision 3: Resume-Safe Checkpointing")
body(
    "Every single step checks if its output file already exists before running. "
    "If yes — it skips. This means if the pipeline crashes at Step 10, you just "
    "restart and it picks up from Step 10. Steps 1–9 are skipped instantly."
)
diagram("""
step_cross_validator():
    if cross_validation_report.json exists:
        print("already done — skipping")
        return  ◀── entire step skipped in 0 seconds
    else:
        run analysis...
        write cross_validation_report.json
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 3. FULL END-TO-END FLOW DIAGRAM
# ══════════════════════════════════════════
h1("3. Full End-to-End Flow — All 15 Steps")

diagram("""
╔═════════════════════════════════════════════════════════════════════╗
║           COMPLETE PIPELINE FLOW (run.py executes this)            ║
╚═════════════════════════════════════════════════════════════════════╝

  $ python run.py --source ./source --output ./results
                              │
                              ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 1 — Layer 1: Source Extraction                                │
│  Type: Pure Python (no AI)   Time: ~1 second                       │
│                                                                     │
│  • Scans source/ folder                                             │
│  • Runs language detectors: PL/SQL extractor + Oracle Forms parser  │
│  • Extracts: tables, packages, triggers, views, procedures          │
│  • Saves structured JSON to results/Source_Extraction/              │
│                                                                     │
│  OUTPUT: Config.json, Database.json, Source_Code.json,              │
│          Extraction_Summary.json, Logs.json                         │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 2 — Scan Once: Build File Cache                               │
│  Type: Pure Python (no AI)   Time: ~0.3 seconds                    │
│                                                                     │
│  • Walks source/ folder                                             │
│  • Reads each of the 42 source files                                │
│  • Stores full content in one JSON: {"filepath": "content", ...}   │
│  • Validates all 42 files are cached                                │
│  • After this step: source/ folder is NEVER read again             │
│                                                                     │
│  OUTPUT: results/file_cache.json (366,098 characters)              │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 0 — Rule Annotator                                            │
│  Type: Claude AI × 6 parallel threads    Time: ~25 minutes         │
│                                                                     │
│  WHY: Before deep analysis, inject labels into the code so every   │
│  downstream AI agent sees pre-identified business rules — not raw  │
│  unlabelled code.                                                   │
│                                                                     │
│  HOW:                                                               │
│  • Reads each PL/SQL file from file_cache.json                      │
│  • Sends to Claude with prompt: "find business rules in this code"  │
│  • Claude adds comment lines above relevant code:                  │
│      -- RULE: description of business rule                         │
│      -- CONSTRAINT: hard-coded limit found                         │
│      -- BUSINESS: workflow logic found                             │
│      -- VALIDATION: input check found                              │
│  • Saves annotated copy to results/annotated_sources/              │
│  • 6 files processed simultaneously (parallel threads)             │
│                                                                     │
│  OUTPUT: results/annotated_sources/ (42 annotated files)           │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 3 — Deep Scan Agent                                           │
│  Type: Claude AI    Time: ~25 minutes                               │
│                                                                     │
│  WHY: Build a master catalogue of EVERYTHING in the system —       │
│  every procedure, every business rule, every data operation —      │
│  before the specialised analysis tracks start.                     │
│                                                                     │
│  HOW:                                                               │
│  • 42 files are too large for one Claude call                      │
│  • Split into 3 chunks of 15 files each                            │
│  • Claude processes each chunk: "catalogue everything in these     │
│    files — procedures, rules, data flows, constraints"             │
│  • SELF-CORRECTION: after each chunk, checks which files were      │
│    covered. If any missed → automatically re-scans missing files   │
│    (up to 3 attempts)                                              │
│  • 3 chunk outputs merged into one master DEEP_SCAN_OUTPUT.md     │
│                                                                     │
│  OUTPUT: results/Scan/Chunk_01/02/03_Output.md                     │
│          results/DEEP_SCAN_OUTPUT.md (master merged document)      │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 3.5 — Implicit Rules Extractor                                │
│  Type: Claude AI    Time: ~16 minutes                               │
│                                                                     │
│  WHY: Business rules hidden in code logic (never written in        │
│  comments) are the most dangerous to miss in a migration.          │
│  This step surfaces them all.                                      │
│                                                                     │
│  HOW: 4 separate Claude passes, each focused on a different        │
│  source type:                                                       │
│                                                                     │
│  Pass 1 — Seed/reference data (2 files):                           │
│    "What business rules are implied by these lookup values?"        │
│    e.g. 6 leave types with specific accrual rates                  │
│                                                                     │
│  Pass 2 — Oracle Forms constraints (6 files):                      │
│    "What validation rules does the UI enforce on this form?"        │
│    e.g. hire date cannot be in the past                            │
│                                                                     │
│  Pass 3 — PL/SQL comment rules (22 files):                         │
│    "What rules are implied by IF/CASE conditions without comments?" │
│    e.g. IF emp_type='CONTRACT' AND tenure<6 THEN eligible='N'      │
│                                                                     │
│  Pass 4 — SQL schema constraints (11 files):                       │
│    "What business rules are encoded in CHECK constraints and FKs?" │
│    e.g. SALARY_BASIS IN ('ANNUAL','HOURLY')                        │
│                                                                     │
│  RESULT: 310 implicit rules extracted                               │
│  OUTPUT: results/implicit_rules.json                               │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
          ┌────────────────────┼────────────────────┐
          │                    │                    │               │
          ▼                    ▼                    ▼               ▼
  ┌──────────────┐   ┌──────────────┐   ┌──────────────┐  ┌──────────────┐
  │   BA TRACK   │   │   DA TRACK   │   │   TA TRACK   │  │   AA TRACK   │
  │ Business     │   │ Data         │   │ Technology   │  │ Application  │
  │ Analysis     │   │ Analysis     │   │ Analysis     │  │ Analysis     │
  │              │   │              │   │              │  │              │
  │ Steps 4,5,6  │   │ Steps 6,7,8  │   │ Steps 8,9,10 │  │ Steps 11,12  │
  │ ~85 min      │   │ ~48 min      │   │ ~86 min      │  │ ~23 min      │
  └──────┬───────┘   └──────┬───────┘   └──────┬───────┘  └──────┬───────┘
         │                  │                  │                  │
         └──────────────────┴──────────────────┴──────────────────┘
                                      │
                     ALL 4 TRACKS FINISH, THEN CONTINUE
                                      │
                                      ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 13 — Cross Validator                                          │
│  Type: Claude AI    Time: ~53 minutes                               │
│                                                                     │
│  WHY: Each track (BA/DA/TA/AA) worked independently. They may      │
│  disagree or know different things. This step compares all 4       │
│  tracks against each other to find gaps and contradictions.        │
│                                                                     │
│  HOW:                                                               │
│  1. Load all 14 output files from BA+DA+TA+AA                      │
│  2. Claude extracts entity lists from each track                   │
│  3. Compare: entity in BA but NOT in DA → GAP                     │
│  4. Compare: BA says X=90, DA says X=180 → CONTRADICTION          │
│  5. For each HIGH/MEDIUM gap: auto-fetch source from cache        │
│     and supplement the missing track's output                      │
│                                                                     │
│  RESULT: 18 gaps found (13 auto-resolved), 7 contradictions found  │
│  OUTPUT: results/cross_validation_report.json                      │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 14 — Foundation Documents                                     │
│  Type: Claude AI    Time: ~60-90 minutes                           │
│                                                                     │
│  WHY: Convert raw analysis outputs into polished, professional     │
│  deliverable documents that a business analyst, architect, or      │
│  developer can actually use.                                       │
│                                                                     │
│  HOW:                                                               │
│  • Reads ALL analysis outputs + cross_validation_report.json       │
│  • Claude synthesises everything and writes 25 documents:          │
│                                                                     │
│  Knowledge Graph (5 files):                                        │
│    ENTERPRISE_KNOWLEDGE_GRAPH.json                                 │
│    CANONICAL_ENTERPRISE_MODEL.md                                   │
│    ARCHITECTURE_INVENTORY.md                                       │
│    TRACEABILITY_MATRIX.md                                          │
│    FORWARD_ENGINEERING_INPUT_MAP.md                                │
│                                                                     │
│  Forward Engineering Docs (20 files):                              │
│    01_BRD.md                  Business Requirements Document       │
│    02_BUSINESS_CAPABILITY_MODEL.md                                 │
│    03_USE_CASE_SPECIFICATION.md                                    │
│    04_BUSINESS_PROCESS_MODEL.md                                    │
│    05_DOMAIN_MODEL.md                                              │
│    06_DATA_DICTIONARY.md                                           │
│    07_DATA_MODEL_SPECIFICATION.md                                  │
│    08_ERD.md                  Entity Relationship Diagram          │
│    09_DATA_FLOW_DIAGRAM.md                                         │
│    10_SERVICE_CATALOG.md                                           │
│    11_API_CONTRACT_SPECIFICATION.md                                │
│    12_TECHNOLOGY_BLUEPRINT.md                                      │
│    13_SECURITY_ARCHITECTURE.md                                     │
│    14_NFR_SPECIFICATION.md                                         │
│    15_FORWARD_ENGINEERING_SPECIFICATION.md                         │
│    16_GENERATION_MANIFEST.json                                     │
│    17_FORWARD_ENGINEERING_READINESS_REPORT.md                      │
│    18_DEPLOYMENT_ARCHITECTURE.md                                   │
│    19_FRONTEND_ARCHITECTURE.md                                     │
│    20_UI_UX_SPECIFICATION.md                                       │
│                                                                     │
│  OUTPUT: results/Foundation_KnowledgeGraph/                        │
│          results/ForwardEngineering_Docs/                          │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│  STEP 15 — Gap Hunter (Self-Healing Loop)                           │
│  Type: Claude AI    Time: ~20-30 minutes                           │
│                                                                     │
│  WHY: Even after all previous steps, something may still be        │
│  missing. This step actively hunts for gaps and fills them.        │
│                                                                     │
│  HOW:                                                               │
│  • Scans ALL output files                                          │
│  • Claude identifies: unanswered questions, missing procedures,    │
│    undocumented tables, thin sections                              │
│  • If gaps found → runs targeted Claude calls to fill each gap     │
│  • Loops until K consecutive rounds find NOTHING new              │
│    (loop stops when the system is satisfied it's complete)         │
│                                                                     │
│  OUTPUT: results/gap_hunter_report.json                            │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
                         ✅ PIPELINE COMPLETE
                      results/ has 50+ output files
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 4. PARALLEL TRACKS DETAIL
# ══════════════════════════════════════════
h1("4. The Four Parallel Analysis Tracks (Steps 4–12)")

body(
    "After Step 3.5, the pipeline launches 4 completely independent analysis tracks in "
    "parallel Python threads. Each track has its own agents, its own questions, and "
    "produces its own outputs. They all read from the same file_cache.json but analyse "
    "different dimensions of the system."
)

h2("4.1 BA Track — Business Analysis (Steps 4, 5, 6)")
diagram("""
BA TRACK
========
Step 4: BA Agent 1 — Structural Scout
    Input:  file_cache.json (all 42 files, Scout requests 41)
    Task:   "Map every business capability to the package that implements it"
    Output: BA_Structural_Scout.md
           • Business capability map
           • Which packages implement which business functions
           • Initial confidence flags on low-certainty findings

Step 5: BA Agent 2 — Deep Analyst (2 turns)
    Turn 1: "Here is the Scout output. List every file you need to deep-analyse"
            Agent requests up to 42 files from cache
    Turn 2: "Here are those files. Extract ALL business rules in detail"
            Agent produces:
            • 140 business rules (BR-01 to BR-140)
            • 7 business domains mapped
            • Process flows for each domain
            • Edge-case pass: "what rules did you miss?"
            • Gap detection: checks if output is complete
    Output: BA_Deep_Analyst.md + BA_Deep_Analyst_Edge.md
""")

h2("4.2 DA Track — Data Analysis (Steps 6, 7, 8)")
diagram("""
DA TRACK
========
Step 6: DA Agent 1 — Data Extractor
    Input:  file_cache.json (Agent requests 40 files focused on schema+data)
    Task:   "Extract the complete data model — every table, column, constraint"
    Output: DA_Data_Extractor.md
           • All 30 tables with columns and data types
           • Foreign key relationships
           • Index definitions
           • Constraint catalogue

Step 7: DA Agent 2 — Data Reviewer (3 passes)
    Pass 1: Deep review of data model quality
           • PII inventory (which columns hold personal data)
           • Migration complexity scoring
           • Redundancy analysis
           • Hidden business rules in schema
    Pass 2: Edge-case pass — "what did Pass 1 miss?"
    Pass 3: Consistency check — 10 cross-checks on data quality
    Gap detection: checks output completeness after each pass
    Output: DA_Data_Reviewer.md + DA_Data_Reviewer_Edge.md
""")

h2("4.3 TA Track — Technology Analysis (Steps 8, 9, 10)")
diagram("""
TA TRACK
========
Step 8: TA Agent 1 — Stack Scout
    Input:  file_cache.json (all 42 files)
    Task:   "Map the complete technology stack, infrastructure, and integrations"
    Output: TA_Stack_Scout.md
           • 13 technology components identified
           • 6 external integrations mapped
           • Security posture overview
           • DevOps maturity assessment
           • Validation queue (items to verify)

Step 9: TA Agent 2 Batch 1 — Deep Analyst
    Input:  First 19 of 38 selected files
    Task:   "Deep-analyse architecture patterns, security details, tech debt"
    Output: TA_Deep_Analyst_Batch1.md

Step 10: TA Agent 2 Batch 2 — Deep Analyst + Synthesis
    Input:  Second 19 files
    Task:   Analyse + synthesis pass (combine Batch1+Batch2)
    Edge-case pass + gap detection
    Gap detection found 18 files Claude wanted but didn't exist in source
    (PKG_DEPARTMENT, HRMS_ADMIN, scheduler_jobs.sql etc.) —
    pipeline correctly reported: "not in cache, cannot supplement"
    Output: TA_Deep_Analyst_Batch2.md + TA_Deep_Analyst_Edge.md
            + TA_Deep_Analyst.md (merged synthesis)
""")

h2("4.4 AA Track — Application Analysis (Steps 11, 12)")
diagram("""
AA TRACK
========
Step 11: AA Agent 1 — App Extractor
    Input:  34 selected files from cache
    Task:   "Build a complete application component map"
    Output: AA_App_Extractor.md
           • Component registry (every module catalogued)
           • Call flow map (which packages call which)
           • Module boundary definitions
           • System inventory

Step 12: AA Agent 2 — Quality Review (3 passes)
    Pass 1: 15 quality findings (QR-001 to QR-015)
           • Architecture violations
           • Risk register entries
           • Integration gaps
    Pass 2: Spot-check of highest-risk areas → 10 more findings
    Pass 3: Edge-case pass → 8 more findings
    Gap detection: "no gaps found — output is complete"
    Total: 33 quality findings, 25 architecture violations, 14 risks

    Also writes structured JSON outputs:
    • component-registry.json
    • call-flow-map.json
    • dependency-graph.json
    • module-boundary-map.json
    • application-risk-register.json
    • architecture-violation-register.json
    • 5 Mermaid diagram files (.mmd) — renderable diagrams
    Output: AA_Quality_Review.md + AA_Quality_Review_Edge.md
            + all JSON/MMD files in D1-application-architecture/
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 5. HOW CLAUDE AI IS USED
# ══════════════════════════════════════════
h1("5. How Claude AI Is Used in the Pipeline")

body(
    "Every step that uses AI follows the same pattern: build a prompt, call Claude, "
    "save the output. The base_runner.py file handles this for every step."
)

h2("5.1 The call_claude() Function")
diagram("""
base_runner.py — call_claude(prompt, label, timeout=1800)
==========================================================

1. Build the Claude CLI command:
   cmd = ["claude", "-p", prompt]

2. Run via Python subprocess:
   result = subprocess.run(cmd, capture_output=True, timeout=timeout)

3. If exit code != 0:
   → Retry up to 5 times with increasing wait (30s, 60s, 120s...)
   → Log each retry attempt

4. If all retries fail:
   → Raise exception → pipeline step FAILS
   → run.py catches it and stops with clear error message

5. Return Claude's text output
   → Each step then parses/saves this output as needed
""")

h2("5.2 Two-Turn Agent Pattern")
body(
    "Most analysis agents use a two-turn pattern. This is important because Claude has "
    "a context window limit — you can't send all 42 files at once. So the agent first "
    "asks which files it needs, then receives only those files."
)
diagram("""
TWO-TURN AGENT PATTERN
======================

TURN 1:
  Prompt to Claude:
    "You are a business analyst. Here is a list of all 42 files.
     Which files do you need to see to do a complete business analysis?
     List only the file paths, one per line."

  Claude responds:
    "ts-plsql.../PKG_EMPLOYEE.pkb
     ts-plsql.../PKG_PAYROLL.pkb
     ts-plsql.../01_core_tables.sql
     ..." (requests up to 42 files)

  Pipeline:
    → Parse Claude's file list
    → Fetch each file content from file_cache.json
    → Build context: {filepath: content, filepath: content, ...}

TURN 2:
  Prompt to Claude:
    "Here are the files you requested:
     [file contents]
     Now do the full business analysis. Extract all business rules,
     process flows, and business capabilities."

  Claude responds:
    [Full analysis document]

  Pipeline:
    → Save to results/Business_Analysis/BA_Deep_Analyst.md

WHY TWO TURNS:
  • Turn 1 lets Claude decide what it needs (agent autonomy)
  • Turn 2 gives Claude exactly what it asked for
  • Avoids sending irrelevant files that waste context
  • Agent can request different files for different tasks
""")

h2("5.3 Multi-Pass Quality Pattern")
body(
    "After an agent produces its main output, it runs additional passes to improve quality:"
)
diagram("""
MULTI-PASS QUALITY PATTERN
===========================

Main pass:    Claude produces initial analysis
              → Save output

Edge-case pass:
              Prompt: "Here is the analysis you just produced.
                       What business rules or edge cases did you miss?
                       Add them now."
              → Claude supplements the output
              → Save edge-case additions

Merge pass:   Prompt: "Combine the main analysis and edge-case additions
                       into one complete, non-redundant document."
              → Final merged output saved

Gap detection:
              Prompt: "Read your merged output. List any files or topics
                       that are referenced but not yet fully covered."
              → If gaps found: fetch missing files from cache, re-analyse
              → Loop until no gaps remain
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 6. STEP-BY-STEP INPUTS AND OUTPUTS
# ══════════════════════════════════════════
h1("6. Every Step — Inputs, What It Does, Outputs")

grid(
    ["Step", "Name", "Type", "Input", "What It Does", "Output"],
    [
        ("1", "Layer 1 Source Extraction", "Pure Python", "source/ folder (42 files)", "Parses all files with language-specific extractors. Extracts DB objects, artifacts, configs.", "Source_Extraction/ (5 JSON files)"),
        ("2", "Scan Once / File Cache", "Pure Python", "source/ folder (42 files)", "Reads every file once. Stores full content in one JSON. Validates completeness.", "file_cache.json (366KB)"),
        ("0", "Rule Annotator", "Claude AI × 6 parallel", "file_cache.json", "Injects -- RULE: -- CONSTRAINT: -- BUSINESS: -- VALIDATION: comments into every PL/SQL file.", "annotated_sources/ (42 files)"),
        ("3", "Deep Scan Agent", "Claude AI (chunked)", "file_cache.json", "Deep-scans all 42 files in 3 chunks. Self-corrects if files missed. Merges into master document.", "Scan/Chunk_01/02/03.md + DEEP_SCAN_OUTPUT.md"),
        ("3.5", "Implicit Rules", "Claude AI (4 passes)", "file_cache.json", "Extracts 310 implicit business rules across 4 passes: seed data, forms, PL/SQL, schema.", "implicit_rules.json"),
        ("4", "BA Agent 1 — Scout", "Claude AI (2-turn)", "file_cache.json", "Maps business capabilities to packages. Builds capability model.", "BA_Structural_Scout.md"),
        ("5", "BA Agent 2 — Deep Analyst", "Claude AI (multi-pass)", "file_cache.json + Scout output", "Extracts 140 business rules across 7 domains. Edge-case and gap-detection passes.", "BA_Deep_Analyst.md + Edge.md"),
        ("6", "DA Agent 1 — Data Extractor", "Claude AI (2-turn)", "file_cache.json", "Extracts full data model: all tables, columns, constraints, relationships.", "DA_Data_Extractor.md"),
        ("7", "DA Agent 2 — Data Reviewer", "Claude AI (3-pass)", "file_cache.json + Extractor output", "Reviews data model quality. PII inventory, migration complexity, redundancy analysis.", "DA_Data_Reviewer.md + Edge.md"),
        ("8", "TA Agent 1 — Stack Scout", "Claude AI (2-turn)", "file_cache.json", "Maps full tech stack, infrastructure, integrations, security posture, DevOps maturity.", "TA_Stack_Scout.md"),
        ("9", "TA Agent 2 Batch 1", "Claude AI", "file_cache.json (19 files)", "Deep tech analysis of first 19 files — architecture patterns, tech debt.", "TA_Deep_Analyst_Batch1.md"),
        ("10", "TA Agent 2 Batch 2 + Synthesis", "Claude AI (multi-pass)", "file_cache.json (19 files) + Batch1", "Deep analysis + synthesis of both batches. Edge-case and gap detection.", "TA_Deep_Analyst_Batch2.md + Edge.md + TA_Deep_Analyst.md"),
        ("11", "AA Agent 1 — App Extractor", "Claude AI (2-turn)", "file_cache.json (34 files)", "Builds component registry, call flow map, module boundaries, system inventory.", "AA_App_Extractor.md + JSON/MMD files"),
        ("12", "AA Agent 2 — Quality Review", "Claude AI (3-pass)", "file_cache.json + Extractor output", "33 quality findings. 25 architecture violations. 14 risk register entries.", "AA_Quality_Review.md + Edge.md"),
        ("13", "Cross Validator", "Claude AI", "All BA+DA+TA+AA outputs", "Finds 18 gaps and 7 contradictions across all 4 tracks. Auto-resolves 13 gaps.", "cross_validation_report.json"),
        ("14", "Foundation Documents", "Claude AI (batched)", "All outputs + cross_validation", "Writes 25 polished deliverable documents: BRD, ERD, API contracts, deployment arch, etc.", "Foundation_KnowledgeGraph/ + ForwardEngineering_Docs/"),
        ("15", "Gap Hunter", "Claude AI (loop)", "All outputs", "Hunts for missing items. Loops until no gaps found. Self-healing.", "gap_hunter_report.json"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 7. HOW RESUME WORKS
# ══════════════════════════════════════════
h1("7. Resume-Safe Pipeline — How Restarts Work")

body(
    "One of the most important features: if the pipeline crashes or times out midway, "
    "you simply run the same command again. It automatically skips all completed steps "
    "and resumes from where it stopped."
)

h2("How Each Step Checks If It's Already Done")
diagram("""
EXAMPLE: Step 13 — Cross Validator

def step_cross_validator(output_dir):

    # CHECK: does output already exist?
    report_path = output_dir / "cross_validation_report.json"

    if report_path.exists() and report_path.stat().st_size > 0:
        print("Cross Validator — already done, skipping.")
        return {"returncode": 0, "duration_s": 0.0}

    # NOT DONE YET — run the analysis
    return _run_or_exit(
        ["python", "cross_validator_runner.py", "--output", str(output_dir)],
        label="Cross Validator",
        timeout=7200
    )

RESULT: If cross_validation_report.json exists → skip in 0 seconds
        If not → run (takes ~53 minutes)
""")

h2("What Happens on Restart")
diagram("""
RESTART SCENARIO: Pipeline crashed during Step 9 (TA Batch 1)

On restart, run.py checks each step in order:

Step 1: Source_Extraction/ exists? YES → skip (0 sec)
Step 2: file_cache.json exists?    YES → skip (0 sec)
Step 0: annotated_sources/ exists? YES → skip (0 sec)
Step 3: DEEP_SCAN_OUTPUT.md?       YES → skip (0 sec)
Step 3.5: implicit_rules.json?     YES → skip (0 sec)
Step 4: BA_Structural_Scout.md?    YES → skip (0 sec)
Step 5: BA_Deep_Analyst.md?        YES → skip (0 sec)
Step 6: DA_Data_Extractor.md?      YES → skip (0 sec)
Step 7: DA_Data_Reviewer.md?       YES → skip (0 sec)
Step 8: TA_Stack_Scout.md?         YES → skip (0 sec)
Step 9: TA_Deep_Analyst_Batch1.md? NO  → RUN (resumes here)
Step 10: ...                       ...

Total time to get back to Step 9: ~1 second
(All previous steps skipped instantly)
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 8. OUTPUT FILES EXPLAINED
# ══════════════════════════════════════════
h1("8. All Output Files — What Each One Contains")

h2("8.1 results/ Folder Structure")
diagram("""
results/
│
├── file_cache.json                  ← All 42 source files stored as JSON
├── implicit_rules.json              ← 310 implicit business rules
├── DEEP_SCAN_OUTPUT.md              ← Master deep scan of all files
├── cross_validation_report.json     ← 18 gaps + 7 contradictions
├── gap_hunter_report.json           ← Final gap hunting results
│
├── Source_Extraction/               ← Step 1 outputs
│   ├── Source_Code.json             ← All parsed code artifacts (1.6 MB)
│   ├── Database.json                ← All DB objects (150 KB)
│   ├── Config.json                  ← Configuration parameters
│   ├── Logs.json                    ← Log file extractions
│   └── Extraction_Summary.json      ← Summary of what was found
│
├── annotated_sources/ (42 files)    ← Step 0 outputs
│   └── [filename].sql/.pkb/.pks     ← Same code + injected comments
│
├── Scan/                            ← Step 3 outputs
│   ├── Chunk_01_Output.md           ← Deep scan of files 1-15
│   ├── Chunk_02_Output.md           ← Deep scan of files 16-30
│   └── Chunk_03_Output.md           ← Deep scan of files 31-42
│
├── Business_Analysis/               ← BA Track outputs
│   ├── BA_Structural_Scout.md       ← Capability map
│   ├── BA_Deep_Analyst.md           ← 140 business rules
│   └── BA_Deep_Analyst_Edge.md      ← Edge case rules
│
├── Data_Analysis/                   ← DA Track outputs
│   ├── DA_Data_Extractor.md         ← Full data model
│   ├── DA_Data_Reviewer.md          ← Quality review, PII, migration
│   └── DA_Data_Reviewer_Edge.md     ← Edge case data findings
│
├── Technology_Analysis/             ← TA Track outputs
│   ├── TA_Stack_Scout.md            ← Full tech stack
│   ├── TA_Deep_Analyst_Batch1.md    ← Deep analysis batch 1
│   ├── TA_Deep_Analyst_Batch2.md    ← Deep analysis batch 2
│   ├── TA_Deep_Analyst_Edge.md      ← Edge cases
│   ├── TA_Deep_Analyst.md           ← Merged synthesis
│   └── ta_agent2_requested_files.json ← Which files were requested
│
├── Application_Analysis/            ← AA Track outputs
│   ├── AA_App_Extractor.md          ← Component map
│   ├── AA_Quality_Review.md         ← 33 quality findings
│   └── AA_Quality_Review_Edge.md    ← Edge case findings
│
├── D1-application-architecture/     ← AA Structured outputs
│   ├── AA_App_Extractor.md
│   ├── application-architecture-summary.md
│   ├── application-interface-catalogue.json
│   ├── application-risk-register.json    ← 14 risks
│   ├── architecture-pattern-report.md
│   ├── architecture-violation-register.json  ← 25 violations
│   ├── call-flow-map.json
│   ├── component-registry.json
│   ├── dependency-graph.json
│   ├── extraction-audit.md
│   ├── forward-engineering-input-map.md
│   ├── module-boundary-map.json
│   ├── open-questions.md
│   ├── strangler-candidate-report.md    ← Migration strategy
│   ├── system-inventory.json
│   └── diagrams/                        ← Mermaid diagram files
│       ├── call-flow-view.mmd
│       ├── component-view.mmd
│       ├── container-view.mmd
│       ├── dependency-view.mmd
│       └── system-context.mmd
│
├── Foundation_KnowledgeGraph/       ← Step 14 outputs (knowledge graph)
│   ├── ENTERPRISE_KNOWLEDGE_GRAPH.json
│   ├── CANONICAL_ENTERPRISE_MODEL.md
│   ├── ARCHITECTURE_INVENTORY.md
│   ├── TRACEABILITY_MATRIX.md
│   └── FORWARD_ENGINEERING_INPUT_MAP.md
│
└── ForwardEngineering_Docs/         ← Step 14 outputs (deliverables)
    ├── 01_BRD.md                    ← Business Requirements Document
    ├── 02_BUSINESS_CAPABILITY_MODEL.md
    ├── 03_USE_CASE_SPECIFICATION.md
    ├── 04_BUSINESS_PROCESS_MODEL.md
    ├── 05_DOMAIN_MODEL.md
    ├── 06_DATA_DICTIONARY.md
    ├── 07_DATA_MODEL_SPECIFICATION.md
    ├── 08_ERD.md
    ├── 09_DATA_FLOW_DIAGRAM.md
    ├── 10_SERVICE_CATALOG.md
    ├── 11_API_CONTRACT_SPECIFICATION.md
    ├── 12_TECHNOLOGY_BLUEPRINT.md
    ├── 13_SECURITY_ARCHITECTURE.md
    ├── 14_NFR_SPECIFICATION.md
    ├── 15_FORWARD_ENGINEERING_SPECIFICATION.md
    ├── 16_GENERATION_MANIFEST.json
    ├── 17_FORWARD_ENGINEERING_READINESS_REPORT.md
    ├── 18_DEPLOYMENT_ARCHITECTURE.md
    ├── 19_FRONTEND_ARCHITECTURE.md
    └── 20_UI_UX_SPECIFICATION.md
""")

doc.add_page_break()

# ══════════════════════════════════════════
# 9. HOW TO RUN
# ══════════════════════════════════════════
h1("9. How to Run the Pipeline")

h2("Prerequisites")
bullet("Python 3.8 or higher installed")
bullet("Claude CLI installed and authenticated (claude -p works in terminal)")
bullet("python-docx installed: pip install python-docx")
bullet("Source files in source/ folder")

h2("Run Command")
diagram("""
cd automated-reverse-engineering-pipeline-main

python run.py --source ./source --output ./results
""")

h2("What Happens")
diagram("""
[STEP 1/15]  Layer 1 — Source Extraction ............... COMPLETE (1.0s)
[STEP 2/15]  Scan Once — Cache All Files ............... COMPLETE (0.3s)
[STEP 0/15]  Rule Annotator ............................ COMPLETE (1487s)
[STEP 3/15]  Deep Scan Agent ........................... COMPLETE (1509s)
[STEP 3.5/15] Implicit Rules .......................... COMPLETE (980s)

PARALLEL ANALYSIS — running BA + DA + TA + AA simultaneously
[STEP 4/15]  BA Agent 1 — Structural Scout  [PARALLEL]
[STEP 6/15]  DA Agent 1 — Data Extractor    [PARALLEL]
[STEP 8/15]  TA Agent 1 — Stack Scout       [PARALLEL]
[STEP 11/15] AA Agent 1 — App Extractor     [PARALLEL]
...
[STEP 13/15] Cross Validator ........................... COMPLETE (3194s)
[STEP 14/15] Foundation Documents ..................... COMPLETE (...)
[STEP 15/15] Gap Hunter ............................... COMPLETE (...)

PIPELINE COMPLETE
""")

h2("If Something Fails")
diagram("""
PIPELINE STOPPED — [STEP 9] TA Agent 2 Batch 1 failed.
Fix the issue and re-run.

Action: Just run the same command again:
  python run.py --source ./source --output ./results

Steps 1, 2, 0, 3, 3.5, 4, 5, 6, 7, 8 will be skipped (already done)
Pipeline resumes from Step 9 automatically.
""")

h2("Estimated Runtime")
grid(
    ["Step", "Type", "Estimated Time"],
    [
        ("Step 1", "Pure Python", "~1 second"),
        ("Step 2", "Pure Python", "~0.3 seconds"),
        ("Step 0", "Claude AI (6 parallel)", "~25 minutes"),
        ("Step 3", "Claude AI", "~25 minutes"),
        ("Step 3.5", "Claude AI (4 passes)", "~16 minutes"),
        ("Steps 4-12", "Claude AI (4 parallel tracks)", "~86 minutes (parallel wall clock)"),
        ("Step 13", "Claude AI", "~53 minutes"),
        ("Step 14", "Claude AI (25 documents)", "~60-90 minutes"),
        ("Step 15", "Claude AI (loop)", "~20-30 minutes"),
        ("TOTAL", "End-to-end", "~5-6 hours"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# 10. GLOSSARY
# ══════════════════════════════════════════
h1("10. Glossary")

grid(
    ["Term", "Meaning"],
    [
        ("Claude AI", "Anthropic's AI model — used for all analysis steps via the claude -p CLI command"),
        ("claude -p", "Claude CLI command that takes a prompt and returns Claude's response as text"),
        ("call_claude()", "Python function in base_runner.py that calls claude -p, handles retries, returns the text"),
        ("file_cache.json", "Single JSON file containing the full text of all 42 source files — created in Step 2, read by all subsequent steps"),
        ("Two-turn agent", "Agent pattern: Turn 1 = ask Claude which files it needs. Turn 2 = give those files and ask for analysis."),
        ("annotated_sources/", "Copies of source files with AI-injected -- RULE:, -- CONSTRAINT:, -- BUSINESS:, -- VALIDATION: comment lines"),
        ("implicit_rules.json", "310 business rules extracted from code logic that were never written in comments"),
        ("Resume-safe", "Pipeline property: every step checks if output exists before running. If yes, it skips. Restart = continue from last failure."),
        ("Parallel tracks", "BA, DA, TA, AA run simultaneously in Python threads — wall clock time = slowest track, not sum of all tracks"),
        ("Edge-case pass", "After main analysis, Claude re-reads its output and asks itself: what did I miss? Adds missing items."),
        ("Gap detection", "After each analysis, Claude checks: are there files or topics I referenced but didn't fully cover? Fetches and fills if found."),
        ("Cross Validator", "Step 13 — compares all 4 tracks against each other to find gaps (one track knows something another doesn't) and contradictions (tracks disagree)"),
        ("Gap Hunter", "Step 15 — self-healing loop that scans all outputs and re-analyses anything still missing, loops until complete"),
        ("Strangler Fig", "Migration strategy: replace old system module by module while it runs, gradually strangling the old system"),
        ("MMD file", "Mermaid diagram file — text-based diagram format renderable in browsers, GitHub, Confluence"),
        ("PL/SQL", "Oracle's procedural programming language — all business logic in the source system"),
        ("Oracle Forms", "Oracle's legacy UI framework — defines screens and validations as XML exports"),
        (".pks / .pkb", ".pks = package specification (interface). .pkb = package body (implementation). Every package has both files."),
        ("BA", "Business Analysis track — extracts business rules, process flows, capabilities"),
        ("DA", "Data Analysis track — extracts data model, ERD, PII inventory, migration complexity"),
        ("TA", "Technology Analysis track — extracts tech stack, security, CI/CD maturity, integrations"),
        ("AA", "Application Analysis track — extracts modules, components, call flows, risk register, architecture violations"),
    ]
)

# ══════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════
from pathlib import Path
out = Path("c:/rev-eng1 test oracle new/automated-reverse-engineering-pipeline-main/automated-reverse-engineering-pipeline-main/ORACLE_HRMS_PIPELINE_EXPLAINED.docx")
doc.save(str(out))
print(f"Saved: {out}")
print(f"Size: {out.stat().st_size:,} bytes")
